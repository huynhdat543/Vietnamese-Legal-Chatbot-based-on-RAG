import pandas as pd
import os
import docx
import re
import json
from datetime import datetime

def get_file_inventory(folder_path):
    docx_files = [file for file in os.listdir(folder_path) if file.endswith(".docx")]
    df_files = pd.DataFrame(docx_files, columns=["File Name"])
    if not df_files.empty:
        df_files["Index"] = df_files["File Name"].str.extract(r"(\d+)").astype(int)
    return df_files

def extract_number_from_filename(filename):
    match = re.match(r"(\d+)\.docx", filename)
    if match:
        return match.group(1)
    return None

def save_to_json(all_data, output_json_path):
    with open(output_json_path, 'w', encoding='utf-8') as json_file:
        json.dump(all_data, json_file, ensure_ascii=False, indent=4)
    print(f"Output path: {output_json_path}")

def process_docx(file_path, all_data, metadata):
    document = docx.Document(file_path)
    stt = extract_number_from_filename(os.path.basename(file_path))
    metadata_row = metadata[metadata['STT'] == int(stt)].iloc[0] \
        if not metadata[metadata['STT'] == int(stt)].empty else None
    
    current_chapter = None
    current_section = None
    current_mini_section = None
    
    chapter_regex = re.compile(r"^(Chương|CHƯƠNG)\s+([IVXLCDM]+|\d+)(\.|:)?")
    section_regex = re.compile(r"^(Mục|MỤC)\s+([IVXLCDM]+|\d+)(\.|:)?")  
    mini_section_regex = re.compile(r"^(Tiểu mục|TIỂU MỤC)\s+([IVXLCDM]+|\d+)(\.|:)?")  
    article_regex = re.compile(r"^Điều\s+(\d+)\.") 

    for i in range(len(document.paragraphs)):
        paragraph_text = document.paragraphs[i].text.strip()

        match_chapter = chapter_regex.match(paragraph_text)
        if match_chapter:
            current_chapter = paragraph_text
            current_section = None
            current_mini_section = None
            if i + 1 < len(document.paragraphs):
                next_paragraph = document.paragraphs[i + 1].text.strip()
                if not section_regex.match(paragraph_text) \
                    or not mini_section_regex.match(paragraph_text) \
                    or not article_regex.match(paragraph_text):
                    current_chapter += f": {next_paragraph}"
            continue

        match_section = section_regex.match(paragraph_text)
        if match_section:
            current_section = paragraph_text
            current_mini_section = None
            continue

        match_mini_section = mini_section_regex.match(paragraph_text)
        if match_mini_section:
            current_mini_section = paragraph_text
            continue

        match_article = article_regex.match(paragraph_text)
        if match_article:
            article = paragraph_text
            content = []
            for j in range(i + 1, len(document.paragraphs)):
                next_paragraph = document.paragraphs[j].text.strip()
                if chapter_regex.match(next_paragraph) or section_regex.match(next_paragraph) \
                        or mini_section_regex.match(next_paragraph) \
                        or article_regex.match(next_paragraph):
                    break
                content.append(next_paragraph)
        
            content_text = "\n".join(content)
            
            all_data.append({
                        "STT": stt,
                        "LoaiVanBan": metadata_row["Loại Văn Bản"],
                        "NoiBanHanh": metadata_row["Nơi Ban Hành"],
                        "SoHieu": metadata_row["Số Hiệu"],
                        "LinhVucNganh": metadata_row["Lĩnh Vực - Ngành"],
                        "NgayBanHanh": metadata_row["Ngày Ban Hành"],
                        "ChuDe": metadata_row["Chủ Đề"],
                        "Chapter": current_chapter,
                        "Section": current_section if current_section else None,
                        "MiniSection": current_mini_section if current_mini_section else None,
                        "Article": article,
                        "OriginalContent": content_text
                    })
    return all_data

def process_folder(folder_path, output_json_path, metadata):
    all_data = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            print(f"Đang xử lý file: {filename}")
            all_data = process_docx(file_path, all_data, metadata)
    save_to_json(all_data, output_json_path)

def create_filter_for_metadata(input_path, output_path):
    try:
        with open(input_path, encoding='utf-8') as file:
            data = json.load(file)
        for item in data:
            ngay_ban_hanh = item.get('NgayBanHanh', '')
            ngay_ban_hanh_formatted = []
            if ngay_ban_hanh:
                try:
                    date_obj = datetime.strptime(ngay_ban_hanh, '%d/%m/%Y')
                    ngay_ban_hanh_formatted = [
                        ngay_ban_hanh,
                        date_obj.strftime('%d-%m-%Y'),
                        date_obj.strftime('%Y/%m/%d'),
                        date_obj.strftime('%Y-%m-%d'),
                        date_obj.strftime('%Y')
                    ]
                except ValueError:
                    print(f"Không thể chuyển đổi ngày '{ngay_ban_hanh}' trong file {input_path}.")
            item['NgayBanHanhFilter'] = ngay_ban_hanh_formatted

            linh_vuc_nganh = item.get('LinhVucNganh', '')
            if ',' in linh_vuc_nganh:
                item['LinhVucNganh'] = [part.strip() for part in linh_vuc_nganh.split(',')]
            else:
                item['LinhVucNganh'] = [linh_vuc_nganh]

            loai_van_ban = item.get('LoaiVanBan', '')
            item['LoaiVanBanFilter'] = loai_van_ban.lower() if loai_van_ban else ''

            noi_ban_hanh = item.get('NoiBanHanh', '')
            item['NoibanHanhFilter'] = noi_ban_hanh.lower() if noi_ban_hanh else ''

        with open(output_path, 'w', encoding='utf-8') as file:
            json.dump(data, file, ensure_ascii=False, indent=4)
        print(f"Đã cập nhật và lưu file: {output_path}")
    except Exception as e:
        print(f"Lỗi khi xử lý file {input_path}: {e}")